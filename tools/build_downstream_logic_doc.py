from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/dajiangdongqu/Documents/质检最新MVP界面")
OUT = ROOT / "outputs" / "汽车销售对话智能平台_事实层之后各层逻辑与高级功能配置指南.docx"

BLUE = "1F5EFF"
BLUE_DARK = "193B75"
BLUE_LIGHT = "EAF2FF"
GREEN = "0F8A72"
GREEN_LIGHT = "E9F7F3"
ORANGE = "C96A00"
ORANGE_LIGHT = "FFF3E2"
RED = "B42318"
RED_LIGHT = "FDECEC"
GRAY = "667085"
GRAY_LIGHT = "F4F6F8"
INK = "172033"
WHITE = "FFFFFF"
FONT = "Hiragino Sans GB"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name=FONT, size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size=10.5, color=INK, bold=False):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold if bold else run.bold)


def add_text(doc, text="", *, bold=False, color=INK, size=10.5, align=None, before=0, after=5, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.18 * level)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = level == 1 and len(doc.paragraphs) > 5
    return p


def add_callout(doc, title, body, fill=BLUE_LIGHT, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 130, 170, 130, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_table(doc, headers, rows, widths=None, font_size=8.8, header_fill=BLUE_DARK):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(header))
        set_run_font(r, size=font_size, bold=True, color=WHITE)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.width = Inches(widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index % 2 == 1:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_flow(doc, labels):
    table = doc.add_table(rows=1, cols=len(labels))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width = 6.5 / len(labels)
    fills = ["E8F0FF", "E9F7F3", "FFF3E2", "EAF2FF", "FCEAF3", "F4F6F8", "E9F7F3"]
    colors = [BLUE, GREEN, ORANGE, BLUE, "C52C75", GRAY, GREEN]
    for idx, label in enumerate(labels):
        cell = table.cell(0, idx)
        cell.width = Inches(width)
        set_cell_shading(cell, fills[idx % len(fills)])
        set_cell_margins(cell, 120, 80, 120, 80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=8.5, bold=True, color=colors[idx % len(colors)])
    return table


def add_logic_block(doc, inputs, config, logic, output, boundary):
    rows = [
        ("上游输入", inputs),
        ("本层配置", config),
        ("判断逻辑", logic),
        ("输出数据", output),
        ("边界/复核", boundary),
    ]
    add_table(doc, ["项目", "说明"], rows, widths=[1.25, 5.25], font_size=9.2, header_fill=BLUE_DARK)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.68)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.32)
section.footer_distance = Inches(0.34)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.2

for name, size, color, before, after in [
    ("Title", 25, INK, 0, 8),
    ("Subtitle", 12, GRAY, 0, 12),
    ("Heading 1", 17, BLUE_DARK, 16, 8),
    ("Heading 2", 13.5, BLUE, 12, 6),
    ("Heading 3", 11.5, GREEN, 8, 4),
]:
    style = styles[name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = name != "Subtitle"
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Number"):
    style = styles[list_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(10.2)

header = section.header
hp = header.paragraphs[0]
hp.text = "基于真实销售对话的可追溯销售执行与客户决策智能平台"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_paragraph_font(hp, size=8.2, color=GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("汽车销售POC默认配置指南  |  ")
set_run_font(run, size=8, color=GRAY)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(82)
p.paragraph_format.space_after = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("事实层之后各层逻辑与\n高级功能配置指南")
set_run_font(r, size=27, bold=True, color=INK)

add_text(doc, "汽车门店销售质检与客户洞察POC", size=14, color=BLUE, bold=True, after=10)
add_text(doc, "核心原则：一次大模型事实抽取，后续全部基于标准事实表、配置规则、策略库和反馈事件计算。", size=11.5, color=GRAY, after=20)

add_callout(
    doc,
    "文档定位",
    "用于产品评审、规则配置、研发实现和客户材料收集。本文给出的是汽车门店POC默认建议值，不是汽车行业统一标准；品牌SOP、车型知识、门店政策和真实样本必须在上线前校准。",
    fill=BLUE_LIGHT,
    accent=BLUE,
)
add_text(doc, "版本：V1.0  |  日期：2026-08-30", size=9.5, color=GRAY, before=18)
doc.add_page_break()

add_heading(doc, "一、结论与设计原则", 1)
add_callout(
    doc,
    "一句话架构",
    "标准事实表是所有下游结果的唯一业务输入。SOP质检、客户洞察、诊断、策略、生成和高级功能不得重新读取整段原文做自由判断；原文只用于证据回溯、人工复核和播放定位。",
    fill=GREEN_LIGHT,
    accent=GREEN,
)

principles = [
    ("一次抽取", "ASR修正文本只进入一次事实层大模型，输出场景、客户事实、销售行为事实和时间戳证据。"),
    ("事实与评价分离", "客户已表达用途，不等于销售完成了用途询问或复述确认；两者必须分列保存。"),
    ("规则可解释", "每个结果必须展示使用了哪些事实、哪条规则、哪版配置、产生了什么输出。"),
    ("策略不重复诊断", "诊断回答“哪里有问题”，策略回答“销售下一步具体做什么”。"),
    ("生成不再判断", "生成层只把事实、诊断和策略渲染成卡片，不重新创造结论。"),
    ("结果与归因分离", "试驾、复店、成交、败单以CRM、POS或人工反馈为准；只能表达关联，不能直接宣称因果。"),
]
add_table(doc, ["原则", "具体含义"], principles, widths=[1.35, 5.15], font_size=9.4)

add_heading(doc, "二、全链路数据流", 1)
add_flow(doc, ["标准事实表", "SOP/洞察", "诊断层", "策略层", "生成层", "反馈层", "业务结果"])
add_text(doc, "人工修正任意事实后，从SOP/洞察开始重新计算所有下游；不再次调用大模型。", bold=True, color=BLUE, before=8)

add_table(
    doc,
    ["层级", "回答的问题", "主要输入", "配置来源", "主要输出"],
    [
        ("SOP质检", "销售是否完成品牌要求动作", "销售行为事实、适用场景", "客户SOP、扣分与适用规则", "动作完成、得分、证据"),
        ("客户洞察", "客户处于什么需求和决策状态", "客户事实、购买/阻塞信号", "标签树、枚举、计算规则", "画像标签、意向、跟进价值"),
        ("诊断层", "销售执行中存在哪些短板", "客户事实+销售行为事实+SOP结果", "问题库、诊断条件", "问题、风险、原因、证据"),
        ("策略层", "针对每个问题下一步怎么做", "诊断结果+客户事实", "策略库、优先级、条件", "行动步骤、时机、渠道、材料"),
        ("生成层", "如何形成业务人员可用的卡片", "事实+诊断+策略", "卡片模板、展示规范", "等级、风险、败单候选、能力、话术卡片"),
        ("反馈层", "人是否认可，行动是否有效", "各层结果+业务结果", "角色、动作、权限、时限", "反馈事件、复核结果、效果关联"),
    ],
    widths=[0.8, 1.35, 1.45, 1.45, 1.45],
    font_size=8.0,
)

add_heading(doc, "三、下游统一输入：标准事实表接口", 1)
add_text(doc, "本文不展开事实层提示词，只定义下游必须使用的输入接口。每条事实必须包含状态、值和证据；未知不能被默认成否。")
add_table(
    doc,
    ["字段组", "核心字段", "状态枚举", "下游用途"],
    [
        ("客户需求事实", "用途、预算/价格范围、核心关注点、决策人、购车周期、竞品", "已明确/部分明确/未明确/不适用", "洞察、需求诊断、车型策略"),
        ("客户异议事实", "异议类型、异议原因、阻碍动作、客户原话", "存在/未发现/待复核", "异议诊断、策略、败单候选"),
        ("购买与阻塞信号", "试驾、询价、金融、置换、现车、交付、留资、下订、拒绝、停止联系", "已出现/未出现/待复核", "意向、跟进价值、客户等级"),
        ("销售需求挖掘行为", "询问用途、预算、周期、决策人、复述确认", "已执行/未发现/待复核", "SOP、需求确认诊断、能力"),
        ("销售讲解与回应", "需求承接、功能说明、价值解释、异议澄清、证据或方案、确认反应", "已执行/未发现/待复核", "讲解诊断、异议诊断、优秀话术"),
        ("销售推进与约定", "试驾邀约、报价、留资、下一动作、时间、渠道、客户确认", "已执行/未发现/待复核", "推进诊断、策略、跟进闭环"),
        ("时序证据", "会话片段、说话人、开始/结束时间、原文、音频位置", "有效/排除/角色待复核", "链路回溯、冲突处理、播放定位"),
    ],
    widths=[1.15, 2.4, 1.25, 1.7],
    font_size=8.2,
)

add_heading(doc, "四、SOP质检层", 1)
add_logic_block(
    doc,
    "适用场景、销售行为事实、客户已表达事实、证据时间戳。",
    "SOP动作库、动作适用条件、必选/可选、权重、缺失扣分、复核要求。",
    "先判断动作是否适用；再判断销售是否执行；最后计算完成度与分数。客户主动表达不等于销售完成确认动作。",
    "SOP动作结果表、质检得分、扣分明细、证据、复核状态。",
    "ASR角色待复核、片段被排除或证据不足时输出“待复核”，不能直接扣分。",
)
add_heading(doc, "4.1 默认判断函数", 2)
add_table(
    doc,
    ["函数", "规则表达", "说明"],
    [
        ("动作适用", "适用场景命中 AND 前置条件成立", "例如客户存在异议时，“处理异议”动作才适用。"),
        ("动作完成", "销售行为事实=已执行 AND 满足最低结构", "不能仅因客户提到预算就判定销售询问预算。"),
        ("动作部分完成", "完成部分步骤但缺少必选步骤", "例如解释了产品但未确认客户是否理解。"),
        ("动作待复核", "关键角色不明 OR 重叠语音 OR 证据冲突", "不进入自动扣分。"),
        ("得分", "基础分 - Σ适用且未完成动作扣分", "关闭规则或不适用动作不扣分。"),
    ],
    widths=[1.25, 2.6, 2.65],
    font_size=8.8,
)
add_heading(doc, "4.2 SOP动作结果数据表", 2)
add_table(
    doc,
    ["字段", "含义", "示例"],
    [
        ("动作编码/名称", "客户可理解名称与系统稳定编码", "询问购车周期"),
        ("适用状态", "适用/不适用/待复核", "适用"),
        ("完成状态", "已完成/部分完成/未完成/待复核", "未完成"),
        ("缺失步骤", "缺少的必选动作结构", "未询问预计用车时间"),
        ("扣分", "规则配置的扣分值", "-5"),
        ("证据", "销售行为证据或客户触发证据", "03:48 客户表达时间未定"),
        ("配置版本", "用于复现结果", "长安门店SOP-2026.08"),
    ],
    widths=[1.2, 2.3, 3.0],
    font_size=8.8,
)

add_heading(doc, "五、客户洞察层", 1)
add_callout(doc, "客户洞察的意义", "事实层回答“客户说了什么、发生了什么”；客户洞察层把多个事实按客户标签树、枚举和计算规则组织成可检索、可运营的客户状态。它不能新增事实，只能对已知事实进行分类、组合和计算。", fill=GREEN_LIGHT, accent=GREEN)
add_logic_block(
    doc,
    "客户需求、异议、购买/阻塞信号、沟通许可、时序证据。",
    "标签树、标签枚举、同义词、互斥规则、意向分值、跟进停止条件。",
    "原子事实先映射为标签，再按权重和冲突规则计算意向与跟进价值；只使用客户原话事实，不使用销售单方陈述。",
    "客户画像标签表、意向分、意向等级、跟进价值、标签证据。",
    "未提及不等于否；互斥标签必须通过事实时间和证据强度解决；重要画像默认可复核。",
)
add_heading(doc, "5.1 POC意向计算建议", 2)
add_table(
    doc,
    ["语义事件", "默认分值", "角色要求", "说明"],
    [
        ("客户明确下订/订金/锁车", "+8", "客户", "同一事件每通只计一次"),
        ("近期购车周期明确", "+3", "客户", "建议默认30天内，可由客户调整"),
        ("客户询价/优惠/金融", "+2", "客户", "需排除否定和转述"),
        ("客户同意试驾", "+2", "客户", "销售单方面邀约不计分"),
        ("客户接受明确下一步", "+2", "客户", "需有动作或时间"),
        ("客户明确需求或预算边界", "+1", "客户", "部分明确也可计基础分"),
        ("客户明确拒绝", "-8", "客户", "进入停止或人工确认"),
        ("客户明确无购买需求", "-6", "客户", "避免把礼貌沟通判为意向"),
    ],
    widths=[2.45, 0.8, 0.8, 2.45],
    font_size=8.7,
)
add_text(doc, "默认阈值：高意向≥7；中高意向≥4；中意向≥1；低于0为低意向；无有效事实为无法判断。上述数值仅为POC初值。", bold=True, color=BLUE)
add_heading(doc, "5.2 客户洞察结果数据表", 2)
add_table(
    doc,
    ["字段", "来源/算法", "输出要求"],
    [
        ("标签组/标签值", "事实字段→标签映射规则", "必须在客户配置枚举内"),
        ("意向分/等级", "购买信号分+阻塞信号分", "展示加减分明细"),
        ("跟进价值", "意向等级+停止条件+沟通许可", "高优先/普通/低优先/暂停"),
        ("证据", "参与计算的客户原话", "时间戳、说话人、音频跳转"),
        ("冲突状态", "同类标签互斥或事实冲突", "无冲突/待复核/人工已确认"),
    ],
    widths=[1.3, 2.6, 2.6],
    font_size=8.8,
)

add_heading(doc, "六、诊断层：销售短板分析", 1)
add_logic_block(
    doc,
    "标准事实表、SOP动作结果、客户洞察标签。",
    "问题库、适用条件、命中条件、风险等级、扣分、证据选择、人工复核规则。",
    "对客户事实和销售动作进行交叉判断，识别销售执行缺口；诊断不否定客户已表达的事实。",
    "问题编码、问题名称、风险、命中原因、已知/部分/缺失、证据、可挽回、复核状态。",
    "高风险、责任认定、角色冲突默认进入人工复核；规则未命中不能自由生成问题。",
)
add_heading(doc, "6.1 关键诊断函数", 2)
add_table(
    doc,
    ["诊断问题", "默认命中逻辑", "友好呈现"],
    [
        ("销售需求确认不足", "客户需求维度存在未明确/部分明确，或销售未主动询问/复述确认", "分行展示：已明确、部分明确、仍未明确、销售未确认；明确说明“不否定客户主动表达”。"),
        ("购车周期未确认", "购车周期=未明确 AND 销售询问周期=未执行", "指出缺少周期信息，不重复要求补问已经明确的其他维度。"),
        ("产品讲解匹配不足", "客户关注点存在 AND 销售需求对应讲解=未执行", "列出客户关注点和未形成的对应关系。"),
        ("异议处理不足", "客户异议存在 AND 澄清/回应/确认结构未达到阈值", "分开列异议原因、已完成回应、仍缺步骤。"),
        ("试驾推进不足", "客户存在体验价值或试驾信号 AND 销售未形成邀约/预约", "说明应验证的关注点，而不是泛化邀约。"),
        ("离店跟进闭环缺失", "下一动作、时间、渠道、客户确认中缺少必选项", "列出已约定和未约定项。"),
    ],
    widths=[1.35, 2.55, 2.6],
    font_size=8.1,
)
add_heading(doc, "6.2 诊断结果数据表", 2)
add_table(
    doc,
    ["字段", "说明"],
    [
        ("诊断编码/问题名称", "稳定编码用于策略绑定；名称供业务人员阅读。"),
        ("问题分类/风险等级", "需求挖掘、产品讲解、体验推进、异议、闭环等；低/中/中高/高。"),
        ("命中规则", "记录本次使用的规则编码和条件表达式。"),
        ("形成依据", "以分行结构展示已明确、部分明确、未明确、销售未完成动作。"),
        ("证据", "事实字段证据，不直接重扫整段原文。"),
        ("可挽回/需复核", "用于策略优先级和人工审核。"),
    ],
    widths=[1.6, 4.9],
    font_size=9.0,
)

add_heading(doc, "七、策略层：针对诊断问题生成销售行动", 1)
add_callout(doc, "与诊断层的区别", "诊断层输出“销售需求确认不足”；策略层必须输出“已知什么、只补什么、什么时候通过什么渠道做、准备什么材料、什么结果算完成”。若没有匹配策略，应明确显示待配置，不能重复诊断文案或自动编造。", fill=ORANGE_LIGHT, accent=ORANGE)
add_logic_block(
    doc,
    "诊断问题、客户已知事实、客户标签、风险等级、现有下一步约定。",
    "策略库、绑定问题、适用条件、优先级、行动步骤、时机、渠道、材料、店长介入。",
    "问题编码精确匹配策略；再用客户状态筛选适用策略；同一目标合并去重；按风险和优先级排序。",
    "策略编码、标题、目标、步骤、时机、渠道、材料、完成标准、店长介入。",
    "未配置策略输出“待配置策略”；策略不改变事实和诊断。",
)
add_heading(doc, "7.1 POC默认策略库", 2)
add_table(
    doc,
    ["诊断问题", "策略目标", "默认行动步骤", "时机/渠道"],
    [
        ("未问候开场", "修复首次体验并获得沟通许可", "首次跟进个性化称呼→说明身份→回顾本次看车→确认是否方便沟通", "当天首次跟进；电话/微信"),
        ("销售需求确认不足", "复用已知事实，只补缺口", "先复述已知→确认部分明确边界→仅补问未明确项→匹配车型方案", "24小时内；电话/微信"),
        ("购车周期未确认", "明确用车关键节点", "确认计划用车时间、牌照/置换/交付节点→约定匹配节点的试驾/报价/复店", "24小时内"),
        ("价格异议处理不足", "确认预算差距并提供落地方案", "确认差距→拆解落地价→提供金融/权益选项→确认可接受边界", "24小时内；落地价清单"),
        ("试驾推进不足", "把关注点转化为体验验证", "选择1～3个关注点→设计试驾体验项→邀约→确认时间", "3天内；试驾路线卡"),
        ("决策链未闭合", "让关键决策人参与", "确认影响人关注点→发送决策材料→邀约共同到店/视频沟通", "3天内"),
        ("竞品对比不足", "建立客观比较标准", "确认比较维度→引用可追溯事实→解释适配差异→邀请体验", "24小时内；对比清单"),
        ("跟进闭环缺失", "形成可执行约定", "总结共识→明确动作→明确时间→明确渠道→获得客户确认", "离店前或当天"),
    ],
    widths=[1.25, 1.35, 2.65, 1.25],
    font_size=7.7,
)
add_heading(doc, "7.2 策略冲突与去重", 2)
add_bullet(doc, "已明确的客户事实不得再次作为“补齐项”；只能复述确认或确认边界。")
add_bullet(doc, "预算未确认、购车周期未确认等专项诊断由专项策略处理，需求补全策略不重复询问。")
add_bullet(doc, "同一客户同一时点最多保留3个主策略；高风险、强可挽回和明确业务时限优先。")
add_bullet(doc, "客户明确拒绝联系时，停止普通跟进策略，仅保留人工复核或合规处理。")

add_heading(doc, "八、生成层：把确定结果组织成业务卡片", 1)
add_logic_block(
    doc,
    "标准事实表、诊断结果、策略结果、生成规范。",
    "卡片类型、字段顺序、模板、语气、禁用表述、是否允许表达润色。",
    "模板填充和排序；不重新判断事实、不重新诊断、不重新生成策略。默认不单独生成“下一步跟进建议”，直接展示策略层行动。",
    "客户等级、败单候选、优秀话术候选、风险提醒、销售能力诊断等卡片。",
    "允许大模型润色时只能改写表达，禁止新增事实、问题、策略或因果结论。",
)
add_heading(doc, "8.1 生成卡片数据表", 2)
add_table(
    doc,
    ["字段", "说明"],
    [
        ("卡片类型/标题", "由模板配置决定，不由模型自由创建。"),
        ("来源对象", "引用事实编码、诊断编码和策略编码。"),
        ("结构化内容", "按“已知—问题—行动—时机—证据”渲染。"),
        ("行动按钮", "采纳、调整、驳回、复核、店长介入等。"),
        ("人工状态", "待确认/已通过/已驳回/修改后通过。"),
        ("版本快照", "记录生成规范和上游配置版本。"),
    ],
    widths=[1.45, 5.05],
    font_size=9.0,
)

add_heading(doc, "九、反馈层：采纳、审核与业务结果闭环", 1)
add_logic_block(
    doc,
    "事实、诊断、策略、卡片、人工修正和CRM/POS业务结果。",
    "反馈角色、可执行动作、作用对象、权限、时限、通知渠道。",
    "反馈以事件追加方式记录；反馈不会直接覆盖事实，事实变更必须生成新的人工修正版本并重算下游。",
    "反馈事件表、复核状态、采纳率、策略有效性、业务结果关联。",
    "成交/败单只能由CRM、POS或有权限的人工确认；不得由会话模型直接认定。",
)
add_table(
    doc,
    ["角色", "默认反馈动作", "作用对象"],
    [
        ("销售", "采纳、不适用、修改后使用、已跟进", "策略、生成卡片"),
        ("店长", "认可、驳回、修改原因、标记可挽回、介入", "诊断、策略、高意向预警"),
        ("质检员", "确认问题、驳回问题、需复核", "诊断、扣分、证据"),
        ("内训师", "话术通过、驳回、优化后再审、加入陪练", "优秀话术、能力短板"),
        ("运营", "成交、败单、复店、策略有效/无效", "业务结果、效果归因"),
    ],
    widths=[1.0, 2.7, 2.8],
    font_size=8.8,
)

add_heading(doc, "十、证据、时序、溯源与冲突处理", 1)
add_table(
    doc,
    ["机制", "默认规则"],
    [
        ("证据定位", "每条事实保留会话ID、片段ID、说话人、开始/结束时间、原文和音频跳转位置。"),
        ("下游溯源", "诊断引用事实编码；策略引用诊断编码；卡片引用策略/诊断/事实编码。"),
        ("时序", "客户触发、销售回应、客户反应必须按时间排序；相邻动作超过配置窗口则拆分事件。"),
        ("角色冲突", "说话人未稳定标定或重叠语音无法区分时，标为角色待复核，不进入高风险扣分和优秀话术。"),
        ("事实冲突", "同一字段多值冲突时优先客户直接表达、时间较新、人工确认；仍冲突则保留多值并待复核。"),
        ("配置冲突", "客户级>品牌级>区域级>行业默认；同级以生效时间和已发布版本为准。"),
        ("人工修正", "产生新事实版本；自动触发SOP、洞察、诊断、策略、生成重算；不再次调用事实层大模型。"),
    ],
    widths=[1.25, 5.25],
    font_size=9.0,
)

add_heading(doc, "十一、高级功能一：败单分析", 1)
add_logic_block(
    doc,
    "客户阻塞信号、异议、销售问题、策略执行、CRM败单状态和客户反馈。",
    "候选原因库、映射条件、主次原因规则、可挽回规则、审核角色。",
    "会话只能生成候选原因；CRM/POS/人工确认败单后，才进入正式败单归因。原因按证据完整性、时间接近度和阻断强度排序。",
    "候选败单原因、正式败单原因、主/次原因、可挽回、证据、审核状态。",
    "没有真实败单结果时禁止输出“已败单”；只能显示“候选原因待确认”。",
)
add_table(
    doc,
    ["候选原因", "默认触发条件", "反例/排除"],
    [
        ("价格/金融不匹配", "价格异议存在+方案未被接受+后续停止推进", "只询问价格但继续试驾不算败单"),
        ("产品/配置不匹配", "明确功能缺口+无可接受替代方案", "销售单方说不适合不能作为客户原因"),
        ("竞品流失", "客户明确偏向竞品+本店无后续推进", "仅提及竞品不等于流失"),
        ("决策链未闭合", "关键决策人未参与+客户无法继续决策", "已约家人共同到店时不成立"),
        ("销售执行问题", "诊断高风险问题+对应策略未执行+客户状态倒退", "不能把所有败单都归因销售"),
        ("时间/外部条件", "客户明确暂缓、指标/交付/用车时间不匹配", "应与无需求区分"),
    ],
    widths=[1.45, 3.1, 1.95],
    font_size=8.5,
)

add_heading(doc, "十二、高级功能二：下一步最佳行动", 1)
add_callout(doc, "层级归属", "下一步最佳行动属于策略层，不再由生成层重复判断。生成层只把命中的策略组织成销售可执行卡片。", fill=ORANGE_LIGHT, accent=ORANGE)
add_table(
    doc,
    ["步骤", "逻辑"],
    [
        ("1. 读取问题", "读取诊断问题、风险和可挽回状态。"),
        ("2. 读取客户状态", "读取已知、部分明确、未明确事实以及沟通许可。"),
        ("3. 匹配策略", "按问题编码、客户条件和品牌策略库匹配。"),
        ("4. 冲突去重", "删除重复补问，专项问题交给专项策略。"),
        ("5. 排序", "按风险、业务时限、客户意向和策略优先级排序。"),
        ("6. 输出", "给出目标、步骤、时机、渠道、材料和完成标准。"),
    ],
    widths=[1.2, 5.3],
    font_size=9.0,
)

add_heading(doc, "十三、高级功能三：销售能力诊断与画像", 1)
add_logic_block(
    doc,
    "多通接待的诊断结果、SOP结果、证据样本和人工反馈。",
    "能力维度、问题映射、统计窗口、最少样本、优势/短板阈值、陪练任务。",
    "单通输出“本次能力表现”；达到最少样本后按维度汇总，区分稳定能力、优势项和短板。",
    "本次能力卡、销售画像、证据样本、推荐陪练场景。",
    "不纳入合规表现、业务结果和改进趋势；成交结果不能直接代表销售能力。",
)
add_table(
    doc,
    ["能力维度", "映射问题示例", "默认形成条件"],
    [
        ("接待与关系建立", "未问候、开场机械、沟通许可缺失", "至少10通有效接待"),
        ("需求洞察", "需求确认不足、周期/决策人未确认", "同类场景覆盖≥5通"),
        ("产品价值表达", "讲解泛化、未结合客户需求", "需有客户需求和讲解证据"),
        ("异议处理", "未澄清原因、未确认异议变化", "需有真实异议样本"),
        ("体验与成交推进", "未试驾推进、报价后无行动", "需有适用场景"),
        ("跟进闭环", "无动作/时间/渠道/客户确认", "需有离店或跟进场景"),
    ],
    widths=[1.55, 3.0, 1.95],
    font_size=8.6,
)
add_text(doc, "核心能力表示可持续、可复用的能力维度；优势项表示该销售相对基线表现更好的具体行为。优势必须有足够样本和正向证据，不能由单通录音直接下结论。", color=BLUE_DARK, bold=True)

add_heading(doc, "十四、高级功能四：优秀话术挖掘", 1)
add_callout(doc, "核心原理", "优秀话术不是“销售说了某个关键词”，而是销售在特定客户场景中完成有效行为，并使客户产生可观察的正向反应或会话状态跃迁。关键词只可用于召回线索，不可作为最终入选依据。", fill=GREEN_LIGHT, accent=GREEN)

add_heading(doc, "14.1 全局窗口", 2)
add_table(
    doc,
    ["配置项", "默认值"],
    [
        ("客户触发范围", "销售回应前3句或60秒"),
        ("销售回应范围", "连续1～5句，最长120秒"),
        ("客户反应范围", "销售回应后3句或90秒"),
        ("单个候选片段最长时间", "180秒"),
        ("每通录音最多输出候选", "3条"),
        ("必须具备的角色链", "客户→销售→客户"),
        ("弱回应处理", "“嗯、好、知道了”单独出现不算有效反应"),
    ],
    widths=[2.6, 3.9],
    font_size=9.0,
)

add_heading(doc, "14.2 场景目标库", 2)
add_table(
    doc,
    ["场景", "目标", "最低有效结果"],
    [
        ("需求挖掘", "获取并确认客户真实需求", "客户新增表达≥2项需求，销售复述确认≥1项"),
        ("产品讲解", "将需求与车型功能、价值连接", "客户确认相关性或主动追问具体功能"),
        ("试驾推进", "把关注点转化为体验项目", "客户同意试驾、预约时间或询问试驾条件"),
        ("价格金融", "明确预算边界并提供可理解方案", "客户同意算价、确认方案或继续讨论支付条件"),
        ("异议处理", "澄清原因并降低阻碍", "异议减弱、接受解释或愿意继续推进"),
        ("竞品对比", "建立客观比较标准", "客户愿意重新比较、体验或进一步了解"),
        ("家人决策", "明确决策角色与关注点", "同意邀请家人、转发方案或安排共同到店"),
        ("成交推进", "形成明确交易动作", "同意报价、锁车、下订或提交资料"),
        ("跟进闭环", "明确下一动作、时间和渠道", "客户明确接受跟进安排"),
    ],
    widths=[1.1, 2.35, 3.05],
    font_size=8.2,
)
add_text(doc, "问候开场属于SOP执行，默认不进入优秀话术库；只有形成可观察的个性化破冰效果时才可作为特殊场景配置。", color=GRAY)

add_heading(doc, "14.3 场景有效行为结构", 2)
add_table(
    doc,
    ["场景", "有效结构", "默认要求"],
    [
        ("需求挖掘", "开放提问→追问→复述确认", "3步至少完成2步，复述确认必选"),
        ("产品讲解", "客户需求→产品功能→客户价值→确认理解", "4步至少完成3步，需求对应必选"),
        ("异议处理", "接纳异议→澄清原因→解释/举证→替代方案→确认", "5步至少完成3步，澄清和确认必选"),
        ("价格金融", "预算边界→费用拆解→多方案比较→确认接受度", "4步至少完成3步"),
        ("竞品对比", "比较维度→客观事实→适配差异→客户确认", "4步至少完成3步"),
        ("试驾推进", "关联关注点→设计体验项→发出邀约→确认时间", "4步至少完成3步"),
        ("跟进闭环", "总结已知→下一动作→时间→渠道→客户确认", "5步至少完成4步"),
    ],
    widths=[1.1, 3.45, 1.95],
    font_size=8.2,
)
add_text(doc, "相邻动作最大间隔默认60秒；超过后视为两段独立沟通。", bold=True, color=BLUE)

add_heading(doc, "14.4 客户有效反应枚举", 2)
add_table(
    doc,
    ["等级", "反应类型", "示例", "是否可入选"],
    [
        ("0", "无效或礼貌回应", "嗯、好、知道了、再看看", "否"),
        ("1", "信息补充", "补充用途、预算、周期、决策人", "基础候选"),
        ("2", "认知接受", "表示听懂、复述方案、继续问细节", "基础候选"),
        ("3", "异议减弱", "接受解释、改变原有否定态度", "强候选"),
        ("4", "行动接受", "同意试驾、算价、报价、留资、复店", "强候选"),
        ("5", "交易动作", "锁车、提交资料、下订", "强候选"),
    ],
    widths=[0.55, 1.4, 3.15, 1.4],
    font_size=8.6,
)

add_heading(doc, "14.5 会话状态跃迁", 2)
add_table(
    doc,
    ["之前状态", "之后状态", "结果"],
    [
        ("需求未知", "需求部分/完全明确", "正向跃迁"),
        ("预算模糊", "价格边界或支付方案明确", "正向跃迁"),
        ("产品不了解", "理解与需求相关的功能价值", "正向跃迁"),
        ("异议强烈", "异议减弱或愿意继续了解", "正向跃迁"),
        ("竞品倾向", "愿意重新比较或体验", "正向跃迁"),
        ("无试驾计划", "愿意试驾或完成预约", "强正向跃迁"),
        ("无下一步", "明确动作、时间和渠道", "强正向跃迁"),
        ("犹豫", "报价、留资、锁车或下订", "强正向跃迁"),
        ("中立/积极", "反感、拒绝或停止沟通", "负向跃迁，淘汰"),
    ],
    widths=[2.25, 2.9, 1.35],
    font_size=8.5,
)

add_heading(doc, "14.6 话术淘汰规则", 2)
for item in [
    "缺少客户触发、销售回应或客户后续反应。",
    "角色待复核、重叠语音无法确定说话人。",
    "只有泛化产品介绍，没有对应客户需求。",
    "连续讲解超过90秒，客户没有参与互动。",
    "客户在后3句内明确拒绝、反感或纠正销售。",
    "使用保证最低价、保证贷款通过等绝对承诺。",
    "贬低竞品、虚构权益、错误介绍配置。",
    "仅靠降价推动，没有需求匹配或价值解释。",
    "客户本来已主动要求试驾或下订，销售没有产生新增价值。",
    "缺少可定位录音、时间戳和原文证据。",
]:
    add_bullet(doc, item)

add_heading(doc, "14.7 产品知识正确性", 2)
add_table(
    doc,
    ["知识类型", "默认校验要求"],
    [
        ("车型配置", "匹配品牌、车型、版本、年款"),
        ("价格与权益", "在有效期、区域和门店适用范围内"),
        ("金融方案", "首付、期限、利率/月供计算误差≤1%"),
        ("库存与交付", "来自当前库存或交付数据，不能凭经验保证"),
        ("售后权益", "匹配品牌正式政策和生效时间"),
        ("竞品信息", "有可追溯知识来源，不允许贬损表达"),
        ("知识版本", "保存版本号、生效时间和来源"),
    ],
    widths=[1.45, 5.05],
    font_size=8.9,
)

add_heading(doc, "14.8 审核与后续效果验证", 2)
add_table(
    doc,
    ["配置项", "默认值"],
    [
        ("门店候选审核人", "店长"),
        ("正式话术库审核人", "店长+内训师"),
        ("审核时限", "48小时"),
        ("正式入库最低审核人数", "2人"),
        ("允许修改次数", "2次"),
        ("审核动作", "通过、驳回、修改后再审"),
        ("高风险/知识冲突", "增加产品专家审核"),
        ("自动入库", "默认关闭"),
    ],
    widths=[2.4, 4.1],
    font_size=8.9,
)
add_table(
    doc,
    ["状态", "默认条件"],
    [
        ("系统候选", "1次完整证据链"),
        ("已审核话术", "店长+内训师审核通过"),
        ("门店验证话术", "同类场景使用≥10次、覆盖≥3名销售"),
        ("标杆话术", "使用≥30次、覆盖≥5名销售，流程推进率较同场景基线提升≥5个百分点"),
        ("降级/停用", "负向反应率≥15%，或出现知识错误、投诉、政策过期"),
    ],
    widths=[1.45, 5.05],
    font_size=8.8,
)
add_text(doc, "效果观察窗口：试驾7天、复店14天、正式报价/金融方案14天、锁车/下订30天、成交或败单60天。效果只能表达为候选关联，不能直接宣称话术导致成交。", bold=True, color=RED)

add_heading(doc, "十五、高级功能五：客户三级等级与店长预警", 1)
add_callout(doc, "改进建议", "POC可保留客户原话词库用于召回，但正式判断应把词语先归一为语义事件，并处理否定、转述、历史经历和销售话语；等级计算只使用客户/主客户原话事实。", fill=BLUE_LIGHT, accent=BLUE)
add_table(
    doc,
    ["等级", "业务含义", "默认语义事件", "建议分值"],
    [
        ("一级", "基础咨询意向", "历史了解、新款/保险/交付地点等常规咨询", "每类+1"),
        ("二级", "深度了解意向", "试驾、等车周期、现车、金融、优惠、算价、交车时间", "每类+3"),
        ("三级", "高意向核心", "订车、置换、报废、近期用车、首付、月供、权益、合适就定", "每类+6"),
    ],
    widths=[0.65, 1.45, 3.55, 0.85],
    font_size=8.4,
)
add_bullet(doc, "同一语义事件每通只计一次；相同词在否定句、假设句、销售话语或转述中不计分。")
add_bullet(doc, "累计达到三级阈值后生成预警卡，并通过企业微信群机器人Webhook推送给店长。")
add_bullet(doc, "预警内容包括会话ID、客户称呼、销售、命中客户原话、累计分、建议介入动作和录音跳转链接。")
add_bullet(doc, "店长反馈“已介入/误报/调整等级”，作为反馈事件记录，不直接修改事实。")

add_heading(doc, "十六、下游结果表总览", 1)
add_table(
    doc,
    ["数据表", "主键", "关键字段", "重算触发"],
    [
        ("SOP动作结果表", "会话ID+动作编码", "适用、完成、扣分、证据", "事实/SOP配置变化"),
        ("客户洞察标签表", "会话ID+标签编码", "标签值、分值、证据、冲突", "客户事实/标签规则变化"),
        ("诊断结果表", "会话ID+问题编码", "风险、原因、规则、证据", "事实/SOP/诊断规则变化"),
        ("策略结果表", "会话ID+策略编码", "目标、步骤、时机、渠道、材料", "诊断/策略库变化"),
        ("生成卡片表", "会话ID+卡片编码", "内容、来源、动作、状态", "上游结果/模板变化"),
        ("反馈事件表", "事件ID", "角色、动作、对象、前后值、时间", "人工操作/业务结果"),
        ("败单分析表", "客户ID+商机ID", "候选原因、正式原因、审核、证据", "业务结果/诊断变化"),
        ("销售能力画像表", "销售ID+周期+能力维度", "样本数、短板、优势、证据样本", "有效会话累计/反馈变化"),
        ("优秀话术候选表", "候选ID", "互动链、状态跃迁、知识校验、审核、效果", "事实/规则/业务结果变化"),
    ],
    widths=[1.35, 1.35, 2.55, 1.25],
    font_size=7.8,
)

add_heading(doc, "十七、重算机制与规则函数", 1)
add_text(doc, "建议将下游计算实现为确定性流水线：")
add_number(doc, "读取当前有效事实版本，构建标准事实表。")
add_number(doc, "运行SOP规则与客户洞察标签规则。")
add_number(doc, "运行诊断问题库，生成诊断结果。")
add_number(doc, "按诊断问题匹配策略库，进行冲突去重和优先级排序。")
add_number(doc, "按生成规范渲染卡片，不新增结论。")
add_number(doc, "保留历史结果版本，新的结果标记当前有效。")
add_number(doc, "反馈事件和CRM/POS结果异步关联，不覆盖原始事实。")

add_table(
    doc,
    ["触发事件", "是否调用大模型", "重算范围"],
    [
        ("ASR文本或角色修正", "是：重新事实抽取", "事实及全部下游"),
        ("人工修改标准事实", "否", "SOP/洞察/诊断/策略/生成"),
        ("SOP或标签规则修改", "否", "对应结果及其下游"),
        ("诊断规则修改", "否", "诊断/策略/生成"),
        ("策略库修改", "否", "策略/生成"),
        ("生成模板修改", "否", "生成卡片"),
        ("CRM/POS/人工业务结果", "否", "反馈、效果验证、败单确认"),
    ],
    widths=[2.2, 1.45, 2.85],
    font_size=8.9,
)

add_heading(doc, "十八、客户需要提供的配置材料", 1)
add_table(
    doc,
    ["材料", "使用层级", "最低内容"],
    [
        ("销售SOP", "SOP质检", "动作、适用场景、必选步骤、扣分、证据要求"),
        ("客户标签树", "客户洞察", "标签组、枚举、互斥、未知值、样例标注"),
        ("问题库与诊断条件", "诊断层", "问题编码、适用条件、命中逻辑、风险、复核规则"),
        ("策略库", "策略层", "绑定问题、适用客户、行动、时机、渠道、材料、完成标准"),
        ("卡片规范", "生成层", "卡片类型、字段顺序、文案规范、按钮、禁用表述"),
        ("反馈角色与动作", "反馈层", "角色、权限、动作、对象、时限、通知方式"),
        ("品牌/车型知识", "知识校验", "车型年款、配置、价格、权益、金融、库存、交付、版本"),
        ("业务结果数据", "效果验证", "CRM/POS字段、客户/销售/车型关联键、时间窗"),
        ("历史优秀/失败样本", "校准", "原文、人工结论、业务结果、错误类型"),
    ],
    widths=[1.4, 1.15, 3.95],
    font_size=8.4,
)

add_heading(doc, "十九、POC校准与验收建议", 1)
add_table(
    doc,
    ["阶段", "建议样本", "验收重点"],
    [
        ("规则冷启动", "每类场景正负样本各20～30通", "适用性、误报、漏报、未知比例"),
        ("小流量试运行", "覆盖≥3家门店、≥10名销售", "角色分离、事实修正率、规则命中一致性"),
        ("优秀话术校准", "每场景候选≥30条", "完整互动链、状态跃迁、知识正确性、人审通过率"),
        ("策略效果", "每策略至少30次触发", "采纳率、执行率、后续状态关联，不直接做因果承诺"),
        ("销售画像", "每销售≥10通有效接待", "样本覆盖、场景分布、证据代表性"),
    ],
    widths=[1.4, 2.05, 3.05],
    font_size=8.6,
)
add_callout(
    doc,
    "最终验收原则",
    "任何结果都应能回答四个问题：用了哪些事实？命中了哪条规则？生成了什么行动？后续由谁确认是否有效？当事实被人工修正时，所有下游结果必须确定性变化，并保留新旧版本。",
    fill=BLUE_LIGHT,
    accent=BLUE,
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
