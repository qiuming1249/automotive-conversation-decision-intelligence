from __future__ import annotations

import html
import json
import sqlite3
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "团队分享_质检产品"
DB_PATH = ROOT / "data" / "poc.sqlite"
DOCX_PATH = OUT / "质检与客户洞察产品思考及Ontology架构_团队分享版.docx"
HTML_PATH = OUT / "质检与客户洞察产品_可交互演示文档.html"

BLUE = "165DFF"
TEAL = "0E8F82"
AMBER = "B7791F"
INK = "172033"
MUTED = "607089"
LINE = "D9E2EF"
LIGHT = "F5F8FC"
FONT = "Microsoft YaHei"


def jloads(value, default):
    try:
        return json.loads(value) if value else default
    except Exception:
        return default


def load_snapshot():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    rows = conn.execute(
        """
        SELECT s.id, s.reception_no, s.store, s.salesperson, s.customer_name,
               s.created_at, s.analysis_status, s.asr_status, s.active_version,
               COUNT(t.id) AS utterances, a.score, a.analyzed_at,
               a.fact_package, a.diagnoses, a.strategies, a.generated_cards
        FROM sessions s
        LEFT JOIN transcripts t ON t.session_id=s.id AND t.version=s.active_version
        LEFT JOIN analyses a ON a.session_id=s.id
        GROUP BY s.id
        ORDER BY s.created_at DESC
        """
    ).fetchall()
    sessions = []
    for row in rows:
        facts = jloads(row["fact_package"], {})
        sessions.append(
            {
                "id": row["id"],
                "receptionNo": row["reception_no"],
                "store": row["store"],
                "salesperson": row["salesperson"],
                "customer": row["customer_name"],
                "createdAt": row["created_at"],
                "analysisStatus": row["analysis_status"],
                "asrStatus": row["asr_status"],
                "utterances": int(row["utterances"] or 0),
                "score": row["score"],
                "analyzedAt": row["analyzed_at"],
                "facts": facts.get("signalFacts") or facts.get("derivedResults") or {},
                "diagnoses": jloads(row["diagnoses"], []),
                "strategies": jloads(row["strategies"], []),
                "cards": jloads(row["generated_cards"], []),
            }
        )
    conn.close()
    scores = [s["score"] for s in sessions if isinstance(s["score"], (int, float))]
    metrics = {
        "sessions": len(sessions),
        "analyzed": len(scores),
        "utterances": sum(s["utterances"] for s in sessions),
        "avgScore": round(sum(scores) / len(scores)) if scores else None,
    }
    preferred = next((s for s in sessions if s["receptionNo"] == "RC-0818-801"), None)
    case = preferred or next((s for s in sessions if s["diagnoses"]), sessions[0] if sessions else {})
    return {"generatedAt": datetime.now().strftime("%Y-%m-%d %H:%M"), "metrics": metrics, "sessions": sessions, "case": case}


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def font_run(run, size=10.5, bold=False, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    font_run(run, 9, color=MUTED)


def add_para(doc, text="", size=10.5, bold=False, color=INK, after=5, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(str(text))
    font_run(r, size, bold, color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.65 + level * 0.45)
    p.paragraph_format.first_line_indent = Cm(-0.28)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    font_run(r, 10.2)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    font_run(r, 16 if level == 1 else 12.5, True, BLUE if level == 1 else INK)
    return p


def add_callout(doc, title, text, tone=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(17.0)
    cell = table.cell(0, 0)
    set_cell_fill(cell, "EEF4FF" if tone == BLUE else "ECF8F5")
    set_cell_margins(cell, 130, 160, 130, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "\n")
    font_run(r, 10.5, True, tone)
    r2 = p.add_run(text)
    font_run(r2, 10.2, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    header_props = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_props.append(repeat_header)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_fill(cell, "EAF0F8")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(header))
        font_run(r, 9.5, True)
    for row in rows:
        cells = table.add_row().cells
        row_props = table.rows[-1]._tr.get_or_add_trPr()
        no_split = OxmlElement("w:cantSplit")
        row_props.append(no_split)
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            font_run(r, 9.2)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def fmt_value(value, limit=120):
    if value is None or value == "":
        return "未提及"
    if isinstance(value, list):
        value = "、".join(fmt_value(v, 50) for v in value[:6])
    elif isinstance(value, dict):
        value = "；".join(f"{k}：{fmt_value(v, 45)}" for k, v in list(value.items())[:6])
    text = str(value).replace("\n", " ")
    return text if len(text) <= limit else text[: limit - 1] + "…"


def build_docx(snapshot):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2

    header = section.header.paragraphs[0]
    header.text = "质检与客户洞察产品｜团队分享材料"
    font_run(header.runs[0], 8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("质检与客户洞察产品思考")
    font_run(r, 25, True, INK)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(18)
    r = p2.add_run("从一次大模型事实抽取，到可配置的诊断、策略、生成与反馈闭环")
    font_run(r, 13, False, BLUE)
    add_callout(doc, "一句话结论", "大模型负责理解非结构化对话并形成有证据的事实；Ontology 负责统一业务语义；规则和策略负责稳定推理；生成层负责把结论转成可执行内容；反馈负责验证效果。")
    add_para(doc, f"分享日期：{snapshot['generatedAt']}　｜　材料来源：原始思考文档、当前POC代码与本地真实分析快照", 9.2, color=MUTED)

    add_heading(doc, "一、我们真正要解决的问题", 1)
    add_table(
        doc,
        ["问题", "表现", "产品判断"],
        [
            ["证据不可信", "ASR漏转、短句与重叠语句角色错、客流边界错", "先做证据修正台，人工版本优先于AI版本"],
            ["模型幻觉", "标签相反、结论缺证据、输出格式漂移", "大模型只抽事实，结论由规则与枚举约束"],
            ["成本过高", "同一通录音为多个标签和功能反复跑全文", "事实层一次完整调用，多层复用"],
            ["客户差异", "品牌、SOP、标签树、问题库与策略都不同", "行业通用事实层 + 客户可配置业务层"],
            ["高级功能难解释", "跟进、败单、话术建议像模型自由发挥", "每个输出都回溯到事实、诊断和策略"],
            ["效果无法闭环", "不知道销售是否采纳、是否复店或成交", "反馈事件连接CRM/POS/人工确认"],
        ],
        [3.0, 6.0, 8.0],
    )

    add_heading(doc, "二、核心产品原则：一次洞察，多层复用", 1)
    add_callout(doc, "固定链路", "录音 → ASR与说话人分离 → 人工修正文本 → 一次事实层抽取 → 诊断规则 → 策略库 → 生成规范 → 反馈事件", TEAL)
    add_table(
        doc,
        ["层级", "输入", "本层职责", "主要配置"],
        [
            ["事实层", "修正后的对话、角色、时间戳", "一次大模型抽取客观事实与原文证据", "全局提示词、7类字段提示词、枚举与证据要求"],
            ["诊断层", "事实包、SOP动作、客户标签", "识别销售短板、风险、扣分与复核要求", "问题库、触发条件、风险等级、扣分规则"],
            ["策略层", "诊断问题、客户事实", "匹配下一步动作、时机、渠道与材料", "策略库、触发问题、优先级、介入条件"],
            ["生成层", "事实、诊断、策略", "按规范形成建议卡、败单候选、话术候选", "卡片模板、必填项、禁用表述、反馈动作"],
            ["反馈层", "诊断、策略、卡片、业务结果", "记录采纳、审核、复店、成交、败单与修正", "反馈角色、动作、作用对象、人工覆盖流程"],
        ],
        [2.2, 4.4, 5.5, 5.0],
    )
    add_para(doc, "口径说明：这里的“一次”是指事实层对完整会话的一次结构化抽取；ASR和角色语义标定属于独立的前置处理，不应混为同一次模型调用。", 9.2, color=MUTED)

    add_heading(doc, "三、Ontology：把销售业务变成机器可理解的世界", 1)
    add_para(doc, "这里的 Ontology 不是哲学讨论，而是对门店销售业务中的实体、属性、关系、事件、状态、规则和证据进行统一定义。它解决的是：不同客户、品牌和功能是否在谈同一种业务对象，以及这些对象之间如何关联。")
    add_table(
        doc,
        ["概念", "回答的问题", "本产品中的例子"],
        [
            ["标签", "这段话提到了什么？", "价格敏感、提及竞品、试驾意愿"],
            ["分类法", "它属于哪一类？", "客户异议 → 价格/产品/竞品/决策/时机"],
            ["本体", "哪些对象存在，如何关联和变化？", "客户—提出—异议；销售—回应—异议"],
            ["知识图谱", "本次真实业务中具体发生了什么？", "临时客户—关注—外放电；销售—解释—放电枪"],
            ["场景—角色—行为", "怎样从录音识别本体实例？", "报价沟通—客户—表达价格异议"],
        ],
        [2.6, 5.0, 9.2],
    )
    add_heading(doc, "3.1 最小业务本体", 2)
    add_table(
        doc,
        ["实体", "关键属性", "关键关系"],
        [
            ["会话/证据", "时间、说话人、文本、是否参检", "证据支持事实或判断"],
            ["客户", "需求、预算、关注点、购买信号", "表达需求、提出异议、接受/拒绝方案"],
            ["销售", "询问、讲解、回应、推进动作", "执行SOP、回应异议、推荐产品"],
            ["产品/车型", "配置、价格、场景、品牌知识", "满足需求、被比较、被推荐"],
            ["异议", "类型、原话、次数、阻碍动作", "被客户提出、被销售回应、进入已解决/未解决状态"],
            ["诊断/策略/卡片", "风险、动作、时机、渠道、状态", "由规则命中、匹配策略、生成业务内容"],
        ],
        [3.0, 6.4, 7.4],
    )
    add_heading(doc, "3.2 场景—角色—行为与本体的关系", 2)
    add_para(doc, "本体规定系统需要认识什么；场景—角色—行为负责从录音中识别出来。例：“这车主要是我老婆开，她还没来试过。”")
    add_table(
        doc,
        ["解析维度", "识别结果", "形成的本体事实"],
        [
            ["场景", "决策确认/离店推进", "当前会话进入决策链判断上下文"],
            ["角色", "客户", "该表达不能误记为销售建议"],
            ["行为", "表达主要使用人；表达关键人员未体验", "妻子—是主要使用人；妻子—尚未参与—试驾事件"],
            ["证据", "原话与时间戳", "录音片段—支持—以上事实"],
        ],
        [2.4, 5.5, 8.9],
    )

    add_heading(doc, "四、事实层：7类最小事实包", 1)
    add_para(doc, "事实层按场景、客户行为和销售行为抽取，不直接输出扣分、策略、意向等级或下一步动作。每类配置独立字段提示词，但合并在一次完整模型请求中输出。")
    add_table(
        doc,
        ["事实类别", "只抽取什么", "后续用途"],
        [
            ["场景事实", "咨询、产品、试驾、报价、金融、竞品等已发生场景", "适用规则与证据定位"],
            ["客户需求与约束", "用途、预算、关注点、决策表达、竞品", "客户洞察、产品匹配"],
            ["客户购买与阻塞信号", "试驾/报价/复店等信号及明确阻塞", "意向规则、跟进优先级"],
            ["客户异议事实", "异议类型、原话、次数、拒绝与阻碍动作", "异议诊断、败单候选"],
            ["销售需求挖掘行为", "销售问了什么、客户怎样回答", "SOP质检、能力诊断"],
            ["销售讲解与异议回应", "需求—讲解对应、回应动作、客户反应", "讲解质量、优秀话术候选"],
            ["销售推进与跟进约定", "已经发生的报价、试驾、资料、时间、渠道和客户回应", "闭环诊断、策略匹配"],
        ],
        [3.2, 8.2, 5.4],
    )
    add_callout(doc, "不输出置信度字段", "大规模运营中不逐条维护模型置信度。可解释性由原文证据、规则命中、风险等级、人工复核和业务结果承担；没有证据时输出“未提及/待复核”。")

    add_heading(doc, "五、SOP质检、客户洞察和诊断如何生成", 1)
    add_table(
        doc,
        ["能力", "上游依赖", "计算逻辑", "客户需提供"],
        [
            ["SOP质检", "销售动作事实、客户回答、场景", "事实动作与SOP要求对照，计算完成/缺失/不适用及扣分", "SOP动作、适用条件、扣分、证据要求"],
            ["客户洞察", "客户需求、信号、异议、证据", "按标签树枚举和规则映射，禁止模型自由造标签", "标签树、枚举、标签规则、样例标注"],
            ["诊断层", "SOP结果、客户标签、事实关系", "问题库规则判断销售短板、风险及是否可挽回", "问题库、诊断条件、风险与复核规则"],
            ["策略层", "诊断问题、客户状态、事实约束", "按问题编码与适用条件匹配策略，按优先级排序", "策略库确认、动作边界、时机、材料"],
        ],
        [2.6, 4.3, 6.1, 4.0],
    )

    add_heading(doc, "六、四个高级功能的底层逻辑", 1)
    add_table(
        doc,
        ["功能", "不是", "正确逻辑", "最终确认"],
        [
            ["下一步跟进建议", "泛泛提醒“加强跟进”", "事实识别阻碍 → 诊断问题 → 策略库匹配动作/时机/材料 → 卡片生成", "销售采纳/不适用/已跟进"],
            ["败单分析", "听一通录音就确定败单原因", "对话只生成候选原因和证据；结合CRM/POS/后续反馈确认真实结果", "店长/运营/CRM结果"],
            ["销售能力诊断", "用一通接待给销售定画像", "单通形成能力表现；多通按同一问题与能力维度聚合后形成画像", "累计样本与人工校准"],
            ["优秀话术挖掘", "模型编一句漂亮话", "完整场景 + 客户问题 + 销售回应 + 客户正向反应 → 候选 → 审核入库", "店长/内训师审核"],
        ],
        [3.0, 4.0, 7.4, 3.1],
    )

    metrics = snapshot["metrics"]
    case = snapshot["case"]
    add_heading(doc, "七、当前POC真实数据快照", 1)
    add_callout(doc, "数据口径", f"截至 {snapshot['generatedAt']}，本地POC库共有 {metrics['sessions']} 条接待记录、{metrics['analyzed']} 条已分析、当前版本转写共 {metrics['utterances']} 句；平均质检分 {metrics['avgScore']}。这些是演示库快照，不代表生产规模或模型准确率。", TEAL)
    add_heading(doc, f"案例：{case.get('receptionNo', '暂无案例')}", 2)
    add_para(doc, f"门店：{case.get('store', '')}　｜　销售：{case.get('salesperson', '')}　｜　转写：{case.get('utterances', 0)}句　｜　质检分：{case.get('score', '待分析')}　｜　状态：{case.get('analysisStatus', '')}", 9.5, color=MUTED)
    fact_rows = []
    for key, value in list(case.get("facts", {}).items())[:7]:
        fact_rows.append([key, fmt_value(value, 180)])
    add_table(doc, ["事实类别", "当前抽取结果（节选）"], fact_rows or [["事实层", "暂无结果"]], [4.2, 12.6])
    diag_rows = []
    for item in case.get("diagnoses", [])[:6]:
        evidence = (item.get("evidence") or [{}])[0]
        diag_rows.append([item.get("issue", ""), item.get("riskLevel", ""), f"{evidence.get('timestamp', '')} {evidence.get('speaker', '')}：{evidence.get('quote', '暂无直接证据')}"])
    add_table(doc, ["诊断问题", "风险", "原文证据（节选）"], diag_rows or [["暂无", "", ""]], [4.1, 2.1, 10.6])
    strategy_rows = [[s.get("issue", ""), s.get("nextBestAction", ""), f"{s.get('timing', '')}｜{s.get('channel', '')}"] for s in case.get("strategies", [])[:6]]
    add_table(doc, ["问题", "匹配策略", "时机与渠道"], strategy_rows or [["暂无", "", ""]], [3.7, 8.9, 4.2])
    add_para(doc, "注意：案例中的质检、洞察和策略均为AI/规则生成并处于需复核状态；真实成交、败单及责任结论必须依赖CRM、POS或人工反馈。", 9.2, color=AMBER)

    add_heading(doc, "八、当前Demo实现了什么", 1)
    for item in [
        "录音上传、阿里云ASR、会话内匿名说话人分离与大模型角色语义标定。",
        "音频回放、转写改字、改角色、拆句、合句、插入漏句、参检控制和人工修正版本。",
        "全局提示词 + 7类事实字段提示词的一次事实层抽取。",
        "销售SOP、客户标签树、诊断规则、策略库、生成规范与反馈角色的可配置界面。",
        "事实、诊断、策略、卡片与原文证据的推理链和语义图谱。",
        "质检与客户洞察结果导出、销售/店长/运营反馈闭环。",
    ]:
        add_bullet(doc, item)
    add_para(doc, "角色分离最新口径：工牌已绑定销售身份；系统只在当前录音内形成匿名说话人簇，用会话内声纹特征细化短句和重叠语句，再由大模型结合整通语义把匿名角色映射为销售/客户。无需外部销售声纹库，也不做身份识别。", 9.6, color=INK)

    add_heading(doc, "九、边界与验收", 1)
    add_table(
        doc,
        ["边界", "原因", "验收方式"],
        [
            ["ASR与角色错误会传导", "错误说话人会污染事实和质检", "角色/文本人工修正率、关键样本准确率"],
            ["事实层仍可能遗漏或误抽", "开放语言无法只靠Schema完全消除幻觉", "字段级准确率、召回率、证据有效率、未知比例"],
            ["规则命中不等于责任认定", "高风险结论需要业务上下文", "高风险必审、普通样本抽检、驳回率"],
            ["会话结论不等于业务结果", "成交/败单发生在录音之外", "CRM/POS/人工结果回流率"],
            ["同一录音重复上传会形成多条记录", "POC用于链路验证，不是生产去重口径", "生产接入时增加音频指纹和业务主键"],
        ],
        [4.0, 6.3, 6.5],
    )

    add_heading(doc, "十、客户配置材料与下一步", 1)
    add_table(
        doc,
        ["客户提供", "作用层", "最小材料"],
        [
            ["销售SOP", "事实/SOP质检", "动作、适用条件、扣分、证据要求"],
            ["客户标签树", "事实/客户洞察", "标签、枚举、互斥关系、未提及口径、样例"],
            ["问题库与诊断条件", "诊断层", "问题名称、触发条件、风险、可挽回、复核要求"],
            ["策略确认", "策略层", "动作、时机、渠道、材料、店长介入边界"],
            ["卡片规范", "生成层", "模板、必填内容、禁用表述、反馈动作"],
            ["真实样本与结果", "校准/反馈", "10-30通代表性ASR、人工标签、CRM/POS结果"],
        ],
        [3.8, 3.0, 9.9],
    )
    add_callout(doc, "建议的下一阶段", "先选一个品牌、一个门店、10-30通覆盖典型场景的录音，建立人工金标准；按ASR/角色、事实字段、SOP、标签、策略四类误差分别校准，再扩展到50-100通验证稳定性和运营成本。")
    add_para(doc, "最终原则：先保证证据可信，再形成事实；先形成事实，再进行推理；先让人工确认建议，再考虑Agent自动执行。", 11, True, BLUE, after=0)

    doc.save(DOCX_PATH)


def build_html(snapshot):
    data_json = json.dumps(snapshot, ensure_ascii=False).replace("</", "<\\/")
    html_text = r'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8" />
<meta name="viewport" content="width=device-width, initial-scale=1" />
<title>质检与客户洞察产品｜团队分享</title>
<style>
:root{--blue:#165dff;--teal:#0e8f82;--amber:#c98212;--coral:#d85b4a;--ink:#172033;--muted:#65748b;--line:#dce4ef;--fill:#f4f7fb;--white:#fff}
*{box-sizing:border-box}html{scroll-behavior:smooth}body{margin:0;color:var(--ink);font-family:"Microsoft YaHei","PingFang SC",Arial,sans-serif;background:#f3f5f8;letter-spacing:0}button,a{font:inherit}button{cursor:pointer}.topbar{position:sticky;top:0;z-index:20;display:flex;align-items:center;justify-content:space-between;gap:16px;padding:12px 28px;background:#fff;border-bottom:1px solid var(--line)}.brand{font-weight:700}.brand small{display:block;margin-top:2px;color:var(--muted);font-weight:400}.topnav{display:flex;gap:4px;overflow:auto}.topnav a{padding:7px 10px;color:#45536a;text-decoration:none;border-bottom:2px solid transparent;white-space:nowrap}.topnav a:hover,.topnav a.active{color:var(--blue);border-color:var(--blue)}.productLink{padding:8px 12px;color:#fff;text-decoration:none;background:var(--blue);border-radius:4px;white-space:nowrap}.shell{max-width:1260px;margin:auto;background:#fff;box-shadow:0 0 0 1px rgba(30,55,90,.05)}section{padding:54px 52px;border-bottom:1px solid var(--line)}.cover{min-height:560px;display:grid;align-content:center;background:#fff}.eyebrow{color:var(--blue);font-weight:700}.cover h1{max-width:900px;margin:16px 0 18px;font-size:44px;line-height:1.2}.lead{max-width:920px;color:#44536a;font-size:19px;line-height:1.8}.meta{display:flex;gap:12px;flex-wrap:wrap;margin-top:28px}.pill{display:inline-flex;align-items:center;padding:6px 10px;border:1px solid #c9d8ef;border-radius:16px;color:#315072;background:#f7faff}.callout{margin-top:28px;padding:18px 20px;border-left:4px solid var(--teal);background:#eef8f6}.sectionHead{display:flex;align-items:end;justify-content:space-between;gap:20px;margin-bottom:28px}.sectionHead h2{margin:0;font-size:28px}.sectionHead p{max-width:620px;margin:0;color:var(--muted);line-height:1.7}.grid{display:grid;gap:14px}.grid3{grid-template-columns:repeat(3,minmax(0,1fr))}.grid4{grid-template-columns:repeat(4,minmax(0,1fr))}.card{padding:18px;border:1px solid var(--line);border-radius:6px;background:#fff}.card h3{margin:0 0 8px;font-size:17px}.card p{margin:0;color:#56647a;line-height:1.65}.problem{border-top:3px solid var(--coral)}.metric strong{display:block;font-size:28px}.metric span{color:var(--muted)}.pipeline{display:grid;grid-template-columns:repeat(5,minmax(0,1fr));gap:8px}.layer{min-height:138px;padding:16px;text-align:left;border:1px solid var(--line);border-radius:5px;background:#fff;transition:.18s}.layer:hover,.layer.active{transform:translateY(-3px);border-color:var(--blue);box-shadow:0 8px 20px rgba(28,70,130,.12)}.layer b{display:block;margin-bottom:8px;color:var(--blue)}.layer span{color:#59677d;line-height:1.6}.detailPanel{margin-top:16px;padding:20px;background:var(--fill);border:1px solid var(--line)}.detailPanel h3{margin:0 0 8px}.ontologyWrap{display:grid;grid-template-columns:minmax(0,1fr) 330px;gap:18px}.graph{min-height:460px;padding:8px;border:1px solid var(--line);background-color:#fbfcfe;background-image:radial-gradient(#dce5f0 1px,transparent 1px);background-size:22px 22px;overflow:auto}.graph svg{width:100%;min-width:760px}.node{cursor:pointer;transform-box:fill-box;transform-origin:center;transition:.18s}.node:hover,.node.active{transform:scale(1.08);filter:drop-shadow(0 7px 8px rgba(20,50,90,.18))}.node rect{stroke-width:1.6}.node text{pointer-events:none}.edge{stroke:#9aabba;stroke-width:1.6;fill:none}.node.active~.edge{stroke:var(--blue)}.node-session rect{fill:#eef5ff;stroke:#3b82f6}.node-person rect{fill:#eef9f6;stroke:#3b9d87}.node-fact rect{fill:#f4f8ff;stroke:#719ed0}.node-reason rect{fill:#fff8e8;stroke:#d7a13a}.node-action rect{fill:#f3f0ff;stroke:#8c76d2}.nodeLabel{font-size:13px;font-weight:700;fill:var(--ink)}.nodeType{font-size:10px;fill:#66758b}.inspector h3{margin-top:0}.inspector ul{padding-left:18px;color:#4c5b70;line-height:1.8}.tabs{display:flex;gap:6px;flex-wrap:wrap;margin-bottom:14px}.tabs button,.sessionBtn{padding:8px 12px;border:1px solid var(--line);border-radius:4px;background:#fff}.tabs button.active,.sessionBtn.active{color:var(--blue);border-color:var(--blue);background:#f2f7ff}.dataLayout{display:grid;grid-template-columns:310px minmax(0,1fr);gap:18px}.sessionList{display:grid;gap:8px;align-content:start}.sessionBtn{text-align:left}.sessionBtn strong,.sessionBtn span{display:block}.sessionBtn span{margin-top:4px;color:var(--muted);font-size:12px}.caseHeader{display:flex;justify-content:space-between;gap:12px;align-items:start}.score{font-size:34px;color:var(--blue);font-weight:700}.caseTabs{display:flex;gap:4px;margin:18px 0;border-bottom:1px solid var(--line);overflow:auto}.caseTabs button{padding:9px 12px;border:0;border-bottom:2px solid transparent;background:transparent;white-space:nowrap}.caseTabs button.active{color:var(--blue);border-color:var(--blue)}.rows{display:grid;gap:9px}.rowBtn{width:100%;padding:13px 14px;text-align:left;border:1px solid var(--line);border-radius:4px;background:#fff}.rowBtn:hover{border-color:#91b4f7;background:#f8fbff}.rowBtn strong{display:block;margin-bottom:5px}.rowBtn span{color:#59677d;line-height:1.6}.risk{float:right;color:var(--amber)}.drawer{position:fixed;inset:0;z-index:40;display:none;background:rgba(20,29,44,.3)}.drawer.open{display:block}.drawerPanel{position:absolute;right:0;top:0;width:min(520px,92vw);height:100%;padding:28px;background:#fff;overflow:auto}.drawerClose{float:right;border:0;background:var(--fill);width:34px;height:34px}.quote{margin:10px 0;padding:13px;border-left:3px solid var(--blue);background:var(--fill);line-height:1.7}.advanced{display:grid;grid-template-columns:240px minmax(0,1fr);gap:18px}.advancedMenu{display:grid;gap:7px;align-content:start}.advancedMenu button{padding:12px;text-align:left;border:1px solid var(--line);background:#fff}.advancedMenu button.active{color:#fff;border-color:var(--blue);background:var(--blue)}.logicFlow{display:flex;align-items:center;gap:8px;flex-wrap:wrap;margin:16px 0}.logicFlow span{padding:8px 10px;background:var(--fill);border:1px solid var(--line)}.logicFlow b{color:var(--blue)}table{width:100%;border-collapse:collapse}th,td{padding:11px 12px;text-align:left;vertical-align:top;border-bottom:1px solid var(--line);line-height:1.55}th{color:#4b5a70;background:#f4f7fb}.note{color:var(--muted);font-size:13px}.footer{padding:34px 52px;color:var(--muted);background:#f8fafc}.toast{position:fixed;right:20px;bottom:20px;padding:10px 14px;color:#fff;background:#26364f;border-radius:4px;opacity:0;transform:translateY(8px);transition:.2s}.toast.show{opacity:1;transform:none}
@media(max-width:900px){.topbar{align-items:flex-start;padding:10px 14px}.topnav{display:none}.productLink{font-size:13px}section{padding:36px 18px}.cover h1{font-size:32px}.grid3,.grid4,.pipeline{grid-template-columns:1fr}.ontologyWrap,.dataLayout,.advanced{grid-template-columns:1fr}.graph{min-height:380px}.sectionHead{align-items:start;flex-direction:column}.sessionList{grid-template-columns:repeat(2,minmax(0,1fr))}.footer{padding:28px 18px}}
</style>
</head>
<body>
<header class="topbar"><div class="brand">质检与客户洞察<small>团队分享演示文档</small></div><nav class="topnav"><a href="#problem">问题</a><a href="#architecture">架构</a><a href="#ontology">Ontology</a><a href="#data">真实数据</a><a href="#advanced">高级功能</a><a href="#delivery">落地</a></nav><a class="productLink" href="http://localhost:5173/" target="_blank">打开当前Demo</a></header>
<main class="shell">
<section class="cover"><div class="eyebrow">一次洞察，多层复用</div><h1>把销售录音从“大模型自由生成”变成可解释、可配置、可运营的业务系统</h1><p class="lead">大模型负责形成有证据的事实，Ontology 统一业务语义，规则与策略控制判断和动作，反馈连接真实业务结果。</p><div class="meta"><span class="pill">事实层一次完整调用</span><span class="pill">五层配置驱动</span><span class="pill">证据可回溯</span><span class="pill">当前POC真实快照</span></div><div class="callout"><b>分享主线：</b>证据可信 → 事实结构化 → 规则诊断 → 策略匹配 → 卡片生成 → 反馈验证</div></section>
<section id="problem"><div class="sectionHead"><h2>01 为什么要重构传统智能质检</h2><p>问题不只是模型准确率，而是证据、成本、客户配置和业务结果没有形成统一闭环。</p></div><div class="grid grid3"><article class="card problem"><h3>证据不稳</h3><p>ASR漏转、角色错、重叠语句和客流边界错误会向所有下游结果传播。</p></article><article class="card problem"><h3>输出漂移</h3><p>标签值矛盾、格式不稳定、没有原文证据，无法直接用于质检与经营。</p></article><article class="card problem"><h3>成本重复</h3><p>几十上百个标签和高级功能反复输入整通录音，Token和时延持续增加。</p></article><article class="card problem"><h3>客户口径不同</h3><p>不同品牌的SOP、标签树、问题库、策略和话术规范均不同。</p></article><article class="card problem"><h3>高级功能难解释</h3><p>跟进、败单和优秀话术如果没有逻辑链，就像模型临场发挥。</p></article><article class="card problem"><h3>效果没有回流</h3><p>不知道建议是否采纳、是否执行、客户是否复店或成交。</p></article></div></section>
<section id="architecture"><div class="sectionHead"><h2>02 五层架构</h2><p>点击每一层查看输入、职责和配置。除事实层外，诊断、策略、生成和反馈均由配置控制。</p></div><div class="pipeline" id="pipeline"></div><div class="detailPanel" id="layerDetail"></div></section>
<section id="ontology"><div class="sectionHead"><h2>03 Ontology 与场景—角色—行为</h2><p>Ontology 规定系统需要认识什么；场景—角色—行为负责从录音中识别并实例化。</p></div><div class="ontologyWrap"><div class="graph"><svg viewBox="0 0 900 460" id="ontologyGraph"></svg></div><aside class="card inspector" id="ontologyDetail"></aside></div><div class="callout"><b>四者区别：</b>标签回答“提到了什么”；分类法回答“属于哪类”；本体回答“对象如何关联和变化”；知识图谱保存“本次真实业务具体发生了什么”。</div></section>
<section id="data"><div class="sectionHead"><h2>04 当前POC真实数据</h2><p>数据来自本地SQLite快照。点击接待记录与结果项查看事实、诊断、策略、卡片和原文证据。</p></div><div class="grid grid4" id="metrics"></div><p class="note">提示：这些是当前演示库接待记录，不代表生产规模、模型准确率或真实经营收益；分析结论默认需复核。</p><div class="dataLayout"><div class="sessionList" id="sessionList"></div><div class="card" id="casePanel"></div></div></section>
<section id="advanced"><div class="sectionHead"><h2>05 四个高级功能如何形成</h2><p>它们不是再跑一次完整录音，而是复用事实、诊断、策略和业务结果。</p></div><div class="advanced"><div class="advancedMenu" id="advancedMenu"></div><div class="card" id="advancedDetail"></div></div></section>
<section id="delivery"><div class="sectionHead"><h2>06 客户要提供什么，怎样落地</h2><p>行业通用事实层保持稳定；SOP、标签、诊断和策略按客户与品牌差异配置。</p></div><table><thead><tr><th>客户材料</th><th>作用层</th><th>最小要求</th></tr></thead><tbody><tr><td>销售SOP</td><td>事实/SOP质检</td><td>动作、适用条件、扣分、证据要求</td></tr><tr><td>客户标签树</td><td>事实/客户洞察</td><td>枚举、互斥关系、未提及口径、样例</td></tr><tr><td>问题库与诊断条件</td><td>诊断层</td><td>触发条件、风险、可挽回、复核要求</td></tr><tr><td>策略确认</td><td>策略层</td><td>动作、时机、渠道、材料、介入边界</td></tr><tr><td>卡片规范</td><td>生成层</td><td>模板、必填项、禁用表述、反馈动作</td></tr><tr><td>真实样本与结果</td><td>校准/反馈</td><td>10-30通代表性ASR、人工标注、CRM/POS结果</td></tr></tbody></table><div class="callout"><b>建议：</b>先选一个品牌、一个门店、10-30通典型录音建立人工金标准；按ASR/角色、事实、SOP/标签、策略分别校准，再扩展到50-100通验证稳定性和运营成本。</div></section>
<footer class="footer"><b>最终原则：</b>先保证证据可信，再形成事实；先形成事实，再进行推理；先让人工确认建议，再考虑Agent自动执行。<br><span id="snapshotTime"></span></footer>
</main>
<div class="drawer" id="drawer"><div class="drawerPanel"><button class="drawerClose" onclick="closeDrawer()">×</button><div id="drawerContent"></div></div></div><div class="toast" id="toast"></div>
<script>
const SNAPSHOT=__DATA__;
const layers=[
 {name:'事实层',tag:'一次大模型',text:'抽取客观事实与原文证据',input:'人工修正后的对话、角色和时间戳',config:'全局系统提示词、7类字段提示词、枚举和证据要求'},
 {name:'诊断层',tag:'规则引擎',text:'识别问题、风险、扣分和复核',input:'事实包、SOP动作、客户标签',config:'问题库、触发条件、风险等级、扣分和复核规则'},
 {name:'策略层',tag:'策略库',text:'匹配动作、时机、渠道和材料',input:'诊断问题、客户事实与约束',config:'触发问题、优先级、动作、材料、店长介入'},
 {name:'生成层',tag:'模板规范',text:'形成建议卡和分析卡',input:'事实、诊断和策略',config:'卡片模板、必填项、禁用表述、是否允许小上下文润色'},
 {name:'反馈层',tag:'事件闭环',text:'记录采纳、审核和业务结果',input:'诊断、策略、卡片与业务结果',config:'反馈角色、动作、作用对象和人工覆盖流程'}];
const pipeline=document.getElementById('pipeline'),layerDetail=document.getElementById('layerDetail');
function selectLayer(i){[...pipeline.children].forEach((x,n)=>x.classList.toggle('active',n===i));const x=layers[i];layerDetail.innerHTML=`<h3>${x.name}｜${x.tag}</h3><p><b>输入：</b>${x.input}</p><p><b>职责：</b>${x.text}</p><p><b>配置：</b>${x.config}</p>`}
layers.forEach((x,i)=>{const b=document.createElement('button');b.className='layer';b.innerHTML=`<b>${i+1}. ${x.name}</b><span>${x.text}</span>`;b.onclick=()=>selectLayer(i);pipeline.appendChild(b)});selectLayer(0);
const ontologyNodes=[
 {id:'session',type:'会话',label:'接待会话',x:90,y:190,tone:'session',desc:'一次可独立分析的销售与客户对话。'},
 {id:'customer',type:'角色',label:'客户',x:260,y:90,tone:'person',desc:'表达需求、异议、购买信号与决策约束。'},
 {id:'sales',type:'角色',label:'销售',x:260,y:285,tone:'person',desc:'执行询问、讲解、回应、报价与推进动作。'},
 {id:'need',type:'事实',label:'客户需求',x:455,y:55,tone:'fact',desc:'用途、预算、关注点、决策与竞品等客观表达。'},
 {id:'objection',type:'事实',label:'客户异议',x:455,y:145,tone:'fact',desc:'异议类型、原话、次数、拒绝与阻碍动作。'},
 {id:'behavior',type:'事实',label:'销售行为',x:455,y:285,tone:'fact',desc:'销售实际问了什么、讲了什么、如何回应和推进。'},
 {id:'diagnosis',type:'推理',label:'诊断问题',x:650,y:145,tone:'reason',desc:'规则基于事实和SOP识别出的短板、风险与复核要求。'},
 {id:'strategy',type:'行动',label:'匹配策略',x:650,y:285,tone:'action',desc:'针对诊断问题配置的动作、时机、渠道与材料。'},
 {id:'evidence',type:'证据',label:'原文片段',x:455,y:400,tone:'session',desc:'带时间戳和说话人的原文，支持事实与判断。'},
 {id:'card',type:'行动',label:'业务卡片',x:820,y:215,tone:'action',desc:'面向销售、店长和运营的建议、分析或候选话术。'}];
const ontologyEdges=[['session','customer','包含'],['session','sales','包含'],['customer','need','表达'],['customer','objection','提出'],['sales','behavior','执行'],['need','diagnosis','参与判断'],['objection','diagnosis','触发'],['behavior','diagnosis','规则判断'],['diagnosis','strategy','匹配'],['strategy','card','生成'],['evidence','need','支持'],['evidence','objection','支持'],['evidence','behavior','支持']];
const svg=document.getElementById('ontologyGraph'),detail=document.getElementById('ontologyDetail');
svg.innerHTML=`<defs><marker id="arr" markerWidth="8" markerHeight="8" refX="7" refY="4" orient="auto"><path d="M0 0L8 4L0 8z" fill="#9aabba"/></marker></defs>`+ontologyEdges.map(e=>{const a=ontologyNodes.find(n=>n.id===e[0]),b=ontologyNodes.find(n=>n.id===e[1]);return `<path class="edge" d="M${a.x+62},${a.y+24} C${(a.x+b.x)/2},${a.y+24} ${(a.x+b.x)/2},${b.y+24} ${b.x-62},${b.y+24}" marker-end="url(#arr)"/><text x="${(a.x+b.x)/2}" y="${(a.y+b.y)/2+10}" text-anchor="middle" font-size="10" fill="#67768b">${e[2]}</text>`}).join('')+ontologyNodes.map(n=>`<g class="node node-${n.tone}" data-id="${n.id}" tabindex="0"><rect x="${n.x-62}" y="${n.y}" width="124" height="48" rx="6"/><text class="nodeType" x="${n.x}" y="${n.y+17}" text-anchor="middle">${n.type}</text><text class="nodeLabel" x="${n.x}" y="${n.y+35}" text-anchor="middle">${n.label}</text></g>`).join('');
function selectOntology(id){document.querySelectorAll('.node').forEach(n=>n.classList.toggle('active',n.dataset.id===id));const n=ontologyNodes.find(x=>x.id===id);const rel=ontologyEdges.filter(e=>e[0]===id||e[1]===id).map(e=>`${ontologyNodes.find(x=>x.id===e[0]).label} ${e[2]} ${ontologyNodes.find(x=>x.id===e[1]).label}`);detail.innerHTML=`<h3>${n.label}</h3><p>${n.desc}</p><h4>关联关系</h4><ul>${rel.map(x=>`<li>${x}</li>`).join('')}</ul>`}
document.querySelectorAll('.node').forEach(n=>{n.onclick=()=>selectOntology(n.dataset.id);n.onkeydown=e=>{if(e.key==='Enter'||e.key===' '){e.preventDefault();selectOntology(n.dataset.id)}}});selectOntology('session');
document.getElementById('metrics').innerHTML=[['接待记录',SNAPSHOT.metrics.sessions],['已分析',SNAPSHOT.metrics.analyzed],['当前版本转写句',SNAPSHOT.metrics.utterances],['平均质检分',SNAPSHOT.metrics.avgScore??'—']].map(x=>`<div class="card metric"><span>${x[0]}</span><strong>${x[1]}</strong></div>`).join('');
let selectedSession=SNAPSHOT.sessions.find(x=>x.receptionNo==='RC-0818-801')||SNAPSHOT.sessions[0],caseTab='facts';
const sessionList=document.getElementById('sessionList'),casePanel=document.getElementById('casePanel');
function renderSessions(){sessionList.innerHTML=SNAPSHOT.sessions.map(s=>`<button class="sessionBtn ${s.id===selectedSession.id?'active':''}" data-id="${s.id}"><strong>${s.receptionNo}</strong><span>${s.store}｜${s.salesperson}</span><span>${s.utterances}句｜${s.score??'待分析'}分</span></button>`).join('');sessionList.querySelectorAll('button').forEach(b=>b.onclick=()=>{selectedSession=SNAPSHOT.sessions.find(s=>s.id===b.dataset.id);caseTab='facts';renderSessions();renderCase()})}
function evidenceHtml(e=[]){return e.length?e.slice(0,4).map(x=>`<div class="quote"><b>${x.timestamp||''} ${x.speaker||''}</b><br>${x.quote||''}</div>`).join(''):'<p class="note">暂无可回溯原文，需人工复核。</p>'}
function openDrawer(title,body){document.getElementById('drawerContent').innerHTML=`<h2>${title}</h2>${body}`;document.getElementById('drawer').classList.add('open')}
window.closeDrawer=()=>document.getElementById('drawer').classList.remove('open');document.getElementById('drawer').onclick=e=>{if(e.target.id==='drawer')closeDrawer()};
function showItem(type,index){const s=selectedSession;let x;if(type==='diagnoses'){x=s.diagnoses[index];openDrawer(x.issue,`<p><b>风险等级：</b>${x.riskLevel||'待配置'}</p><p><b>命中原因：</b>${x.reason||'规则命中'}</p><h3>原文证据</h3>${evidenceHtml(x.evidence)}`)}else if(type==='strategies'){x=s.strategies[index];openDrawer(x.issue,`<p><b>下一步动作：</b>${x.nextBestAction}</p><p><b>时机：</b>${x.timing}</p><p><b>渠道：</b>${x.channel}</p><p><b>材料：</b>${(x.materials||[]).join('、')||'无'}</p>${evidenceHtml(x.evidenceToShow)}`)}else{x=s.cards[index];openDrawer(x.title,`<p><b>状态：</b>${x.status}</p><p>${x.content}</p><h3>原文证据</h3>${evidenceHtml(x.evidence)}`)}}window.showItem=showItem;
function renderCase(){const s=selectedSession;const tabs=[['facts','事实层'],['diagnoses','诊断层'],['strategies','策略层'],['cards','生成卡片']];let rows='';if(caseTab==='facts'){const entries=Object.entries(s.facts||{});rows=entries.length?entries.map(([k,v])=>`<button class="rowBtn" onclick='openDrawer(${JSON.stringify(k)},`+JSON.stringify(`<pre style="white-space:pre-wrap;line-height:1.7">${JSON.stringify(v,null,2)}</pre>`)+`)'><strong>${k}</strong><span>${typeof v==='string'?v:JSON.stringify(v).slice(0,150)}${JSON.stringify(v).length>150?'…':''}</span></button>`).join(''):'<p>暂无事实层结果</p>'}if(caseTab==='diagnoses')rows=s.diagnoses.map((x,i)=>`<button class="rowBtn" onclick="showItem('diagnoses',${i})"><strong>${x.issue}<em class="risk">${x.riskLevel||''}</em></strong><span>${x.reason||'点击查看规则与证据'}</span></button>`).join('')||'<p>暂无诊断</p>';if(caseTab==='strategies')rows=s.strategies.map((x,i)=>`<button class="rowBtn" onclick="showItem('strategies',${i})"><strong>${x.issue}</strong><span>${x.nextBestAction}</span></button>`).join('')||'<p>暂无策略</p>';if(caseTab==='cards')rows=s.cards.map((x,i)=>`<button class="rowBtn" onclick="showItem('cards',${i})"><strong>${x.type}｜${x.title}</strong><span>${x.status}</span></button>`).join('')||'<p>暂无卡片</p>';casePanel.innerHTML=`<div class="caseHeader"><div><h3>${s.receptionNo}</h3><p class="note">${s.store}｜${s.salesperson}｜${s.analysisStatus}</p></div><div class="score">${s.score??'—'}<small style="font-size:12px;color:#65748b">分</small></div></div><div class="caseTabs">${tabs.map(t=>`<button data-tab="${t[0]}" class="${caseTab===t[0]?'active':''}">${t[1]}</button>`).join('')}</div><div class="rows">${rows}</div>`;casePanel.querySelectorAll('.caseTabs button').forEach(b=>b.onclick=()=>{caseTab=b.dataset.tab;renderCase()})}
renderSessions();renderCase();
const advanced=[
 {name:'下一步跟进建议',desc:'告诉销售为什么联系、何时联系、说什么、准备什么，以及下一步要达成什么。',flow:['客户事实与阻碍','诊断问题','策略库匹配','建议卡片','销售采纳反馈'],boundary:'只输出1-3个可执行动作；没有策略时显示待配置，不自动编造。'},
 {name:'败单分析',desc:'输出证据化、多因素、可校正的候选原因，而不是从一通录音确定真实败单。',flow:['对话事实','诊断候选原因','CRM/POS/跟进结果','人工确认','挽回或复盘'],boundary:'没有真实业务结果时必须标记“候选原因”，不能输出确定败单。'},
 {name:'销售能力诊断',desc:'单通形成能力表现，多通接待按统一能力维度聚合后才形成销售画像。',flow:['销售行为事实','诊断问题','能力维度映射','多次接待聚合','训练与复检'],boundary:'不使用一通录音给销售定性；当前Demo设置累计样本门槛。'},
 {name:'优秀话术挖掘',desc:'发现真实接待中有效的表达链，而不是让模型生成一句听起来漂亮的话。',flow:['客户场景与问题','销售完整回应','客户正向反应','候选话术','店长/内训师审核'],boundary:'没有客户反应或证据链不完整时不进入候选池。'}];
let advancedIndex=0;const advMenu=document.getElementById('advancedMenu'),advDetail=document.getElementById('advancedDetail');function renderAdvanced(){advMenu.innerHTML=advanced.map((x,i)=>`<button class="${i===advancedIndex?'active':''}" data-i="${i}">${x.name}</button>`).join('');advMenu.querySelectorAll('button').forEach(b=>b.onclick=()=>{advancedIndex=+b.dataset.i;renderAdvanced()});const x=advanced[advancedIndex];advDetail.innerHTML=`<h3>${x.name}</h3><p>${x.desc}</p><div class="logicFlow">${x.flow.map((f,i)=>`${i?'<b>→</b>':''}<span>${f}</span>`).join('')}</div><p><b>边界：</b>${x.boundary}</p>`}renderAdvanced();
document.getElementById('snapshotTime').textContent=`数据快照生成时间：${SNAPSHOT.generatedAt}`;
const sections=[...document.querySelectorAll('main section[id]')],navs=[...document.querySelectorAll('.topnav a')];addEventListener('scroll',()=>{const y=scrollY+130;let id='';sections.forEach(s=>{if(s.offsetTop<=y)id=s.id});navs.forEach(a=>a.classList.toggle('active',a.getAttribute('href')==='#'+id))},{passive:true});
</script>
</body></html>'''.replace("__DATA__", data_json)
    HTML_PATH.write_text(html_text, encoding="utf-8")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    snapshot = load_snapshot()
    build_docx(snapshot)
    build_html(snapshot)
    print(json.dumps({"docx": str(DOCX_PATH), "html": str(HTML_PATH), "metrics": snapshot["metrics"], "case": snapshot["case"].get("receptionNo")}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
